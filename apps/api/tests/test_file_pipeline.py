from __future__ import annotations

from collections import defaultdict
from pathlib import Path

from docx import Document as DocxDocument

from app.models.workflow import Workflow, WorkflowRun
from app.schemas.execution import RuntimeNodeInput
from app.services.execution import ExecutionContext, execute_file_upload, execute_text_extraction, resolve_runtime_file_path
from app.services.files import DEFAULT_MAX_UPLOAD_SIZE_BYTES, persist_uploaded_file, validate_upload_file
from app.services.parsers import JsonDocumentParser, TxtDocumentParser, get_document_parser, parse_uploaded_file


def create_workflow_and_run(db_session):
    workflow = Workflow(
        name="Upload Workflow",
        description="test workflow",
        status="draft",
        current_version=1,
        latest_saved_version=0,
        graph_json={"id": "g1", "name": "g1", "version": 1, "nodes": [], "edges": []},
    )
    db_session.add(workflow)
    db_session.flush()

    workflow_run = WorkflowRun(
        workflow_id=workflow.id,
        status="running",
        trigger_mode="manual",
        graph_version=1,
        graph_snapshot=workflow.graph_json,
        input_payload={},
        output_payload={},
        preview_payload={},
        log_messages=[],
    )
    db_session.add(workflow_run)
    db_session.flush()
    return workflow, workflow_run


def test_validate_upload_file_accepts_supported_extensions(tmp_path: Path):
    file_path = tmp_path / "sample.txt"
    file_path.write_text("hello world", encoding="utf-8")

    extension, size_bytes = validate_upload_file(file_path, max_size_bytes=DEFAULT_MAX_UPLOAD_SIZE_BYTES)

    assert extension == ".txt"
    assert size_bytes == len("hello world")


def test_validate_upload_file_rejects_oversized_file(tmp_path: Path):
    file_path = tmp_path / "sample.json"
    file_path.write_text('{"hello":"world"}', encoding="utf-8")

    try:
        validate_upload_file(file_path, max_size_bytes=4)
    except ValueError as exc:
        assert "exceeds max upload size" in str(exc)
    else:
        raise AssertionError("Expected oversized file validation error")


def test_persist_uploaded_file_stores_metadata(db_session, tmp_path: Path):
    workflow, workflow_run = create_workflow_and_run(db_session)
    source_path = tmp_path / "sample.csv"
    source_path.write_text("name,score\nada,10\n", encoding="utf-8")

    uploaded_file = persist_uploaded_file(
        db_session,
        source_path=source_path,
        workflow=workflow,
        workflow_run=workflow_run,
        node_id="file-upload-1",
        max_size_bytes=DEFAULT_MAX_UPLOAD_SIZE_BYTES,
    )

    assert uploaded_file.id is not None
    assert uploaded_file.original_name == "sample.csv"
    assert uploaded_file.extension == ".csv"
    assert Path(uploaded_file.storage_path).exists()
    assert uploaded_file.metadata_json["source_path"].endswith("sample.csv")


def test_parser_registry_selects_expected_parser():
    assert isinstance(get_document_parser(".txt"), TxtDocumentParser)
    assert isinstance(get_document_parser(".json"), JsonDocumentParser)


def test_parse_uploaded_file_normalizes_docx(db_session, tmp_path: Path):
    workflow, workflow_run = create_workflow_and_run(db_session)
    source_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Quarterly status report")
    document.add_paragraph("Everything is local-first.")
    document.save(source_path)

    uploaded_file = persist_uploaded_file(
        db_session,
        source_path=source_path,
        workflow=workflow,
        workflow_run=workflow_run,
        node_id="file-upload-1",
        max_size_bytes=DEFAULT_MAX_UPLOAD_SIZE_BYTES,
    )

    parsed = parse_uploaded_file(uploaded_file)

    assert "Quarterly status report" in parsed.text
    assert parsed.metadata["parser"] == "docx"
    assert parsed.metadata["paragraph_count"] == 2


def test_file_upload_and_text_extraction_execution(db_session, tmp_path: Path):
    workflow, workflow_run = create_workflow_and_run(db_session)
    source_path = tmp_path / "notes.txt"
    source_path.write_text("alpha\nbeta\ngamma", encoding="utf-8")

    file_upload_node = {
        "id": "file-upload-1",
        "data": {
            "label": "File Upload",
            "blockType": "file_upload",
            "config": {"accept": ".pdf,.docx,.txt,.csv,.json", "multiple": False, "maxSizeMb": 5},
        },
    }
    text_extraction_node = {
        "id": "text-extract-1",
        "data": {
            "label": "Text Extraction",
            "blockType": "text_extraction",
            "config": {"strategy": "auto"},
        },
    }

    context = ExecutionContext(
        session=db_session,
        workflow=workflow,
        workflow_run=workflow_run,
        runtime_inputs={"file-upload-1": RuntimeNodeInput(files=[str(source_path)])},
        node_inputs={"text-extract-1": {"file": []}},
        node_outputs={},
        memory_store={},
        run_logs=[],
        uploaded_files_by_node=defaultdict(list),
    )

    upload_result = execute_file_upload(file_upload_node, context)
    context.node_inputs["text-extract-1"]["file"].append(upload_result.outputs["file"])

    extract_result = execute_text_extraction(text_extraction_node, context)
    document_payload = extract_result.outputs["document"].value

    assert document_payload["metadata"]["document_count"] == 1
    assert document_payload["documents"][0]["metadata"]["parser"] == "txt"
    assert "alpha" in document_payload["combined_text"]


def test_file_upload_uses_default_local_paths_when_runtime_is_empty(db_session):
    workflow, workflow_run = create_workflow_and_run(db_session)
    sample_path = resolve_runtime_file_path("storage/sample_company_policy.txt")
    sample_path.parent.mkdir(parents=True, exist_ok=True)
    sample_path.write_text("Sample policy content", encoding="utf-8")
    assert sample_path.exists()

    file_upload_node = {
        "id": "file-upload-1",
        "data": {
            "label": "File Upload",
            "blockType": "file_upload",
            "config": {
                "accept": ".pdf,.docx,.txt,.csv,.json",
                "multiple": False,
                "maxSizeMb": 5,
                "defaultLocalPaths": "storage/sample_company_policy.txt",
            },
        },
    }
    context = ExecutionContext(
        session=db_session,
        workflow=workflow,
        workflow_run=workflow_run,
        runtime_inputs={},
        node_inputs={},
        node_outputs={},
        memory_store={},
        run_logs=[],
        uploaded_files_by_node=defaultdict(list),
    )

    upload_result = execute_file_upload(file_upload_node, context)

    files = upload_result.outputs["file"].value
    assert len(files) == 1
    assert files[0]["original_name"] == "sample_company_policy.txt"


def test_file_upload_runtime_file_overrides_library_and_default_paths(db_session, tmp_path: Path):
    workflow, workflow_run = create_workflow_and_run(db_session)
    runtime_path = tmp_path / "runtime-choice.txt"
    runtime_path.write_text("Runtime selected document", encoding="utf-8")
    library_path = tmp_path / "library-choice.txt"
    library_path.write_text("Library default document", encoding="utf-8")
    default_path = resolve_runtime_file_path("storage/default-choice.txt")
    default_path.parent.mkdir(parents=True, exist_ok=True)
    default_path.write_text("Default fallback document", encoding="utf-8")

    library_file = persist_uploaded_file(
        db_session,
        source_path=library_path,
        workflow=workflow,
        workflow_run=workflow_run,
        node_id="library",
        max_size_bytes=DEFAULT_MAX_UPLOAD_SIZE_BYTES,
    )

    file_upload_node = {
        "id": "file-upload-1",
        "data": {
            "label": "File Upload",
            "blockType": "file_upload",
            "config": {
                "sourceMode": "runtime_or_library",
                "libraryFileIds": str(library_file.id),
                "accept": ".txt",
                "multiple": False,
                "maxSizeMb": 5,
                "defaultLocalPaths": "storage/default-choice.txt",
            },
        },
    }
    context = ExecutionContext(
        session=db_session,
        workflow=workflow,
        workflow_run=workflow_run,
        runtime_inputs={"file-upload-1": RuntimeNodeInput(files=[str(runtime_path)])},
        node_inputs={},
        node_outputs={},
        memory_store={},
        run_logs=[],
        uploaded_files_by_node=defaultdict(list),
    )

    upload_result = execute_file_upload(file_upload_node, context)

    files = upload_result.outputs["file"].value
    assert len(files) == 1
    assert files[0]["original_name"] == "runtime-choice.txt"
