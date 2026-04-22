from __future__ import annotations

import csv
import json
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.models.workflow import UploadedFile
from app.services.ocr import run_tesseract_ocr


@dataclass
class ParsedDocument:
    text: str
    metadata: dict[str, Any]


class DocumentParser(ABC):
    extensions: tuple[str, ...] = ()
    parser_name: str = "base"

    def supports(self, extension: str) -> bool:
        return extension.lower() in self.extensions

    @abstractmethod
    def parse(self, uploaded_file: UploadedFile) -> ParsedDocument:
        raise NotImplementedError


class TxtDocumentParser(DocumentParser):
    extensions = (".txt",)
    parser_name = "txt"

    def parse(self, uploaded_file: UploadedFile) -> ParsedDocument:
        path = Path(uploaded_file.storage_path)
        text = path.read_text(encoding="utf-8", errors="ignore")
        return ParsedDocument(text=text, metadata={"parser": self.parser_name, "line_count": len(text.splitlines()), "word_count": len(text.split())})


class CsvDocumentParser(DocumentParser):
    extensions = (".csv",)
    parser_name = "csv"

    def parse(self, uploaded_file: UploadedFile) -> ParsedDocument:
        path = Path(uploaded_file.storage_path)
        rows: list[list[str]] = []
        with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
            reader = csv.reader(handle)
            for row in reader:
                rows.append(row)
        text = "\n".join(", ".join(cell for cell in row) for row in rows)
        return ParsedDocument(text=text, metadata={"parser": self.parser_name, "row_count": len(rows), "detected_tables": 1 if rows else 0, "word_count": len(text.split())})


class JsonDocumentParser(DocumentParser):
    extensions = (".json",)
    parser_name = "json"

    def parse(self, uploaded_file: UploadedFile) -> ParsedDocument:
        path = Path(uploaded_file.storage_path)
        payload = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
        text = json.dumps(payload, indent=2, ensure_ascii=True)
        return ParsedDocument(text=text, metadata={"parser": self.parser_name, "root_type": type(payload).__name__, "word_count": len(text.split())})


class DocxDocumentParser(DocumentParser):
    extensions = (".docx",)
    parser_name = "docx"

    def parse(self, uploaded_file: UploadedFile) -> ParsedDocument:
        document = DocxDocument(uploaded_file.storage_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        text = "\n".join(paragraphs)
        return ParsedDocument(text=text, metadata={"parser": self.parser_name, "paragraph_count": len(paragraphs), "table_count": len(document.tables), "word_count": len(text.split())})


class PdfDocumentParser(DocumentParser):
    extensions = (".pdf",)
    parser_name = "pdf"

    def parse(self, uploaded_file: UploadedFile) -> ParsedDocument:
        reader = PdfReader(uploaded_file.storage_path)
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        text = "\n\n".join(page for page in pages if page)
        failed_pages = sum(1 for page in pages if not page)
        return ParsedDocument(text=text, metadata={"parser": self.parser_name, "page_count": len(reader.pages), "failed_pages": failed_pages, "word_count": len(text.split())})


class OcrDocumentParser(DocumentParser):
    parser_name = "ocr"
    extensions = (".pdf", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp")

    def parse(self, uploaded_file: UploadedFile) -> ParsedDocument:
        language = str((uploaded_file.metadata_json or {}).get("ocr_language") or "eng")
        result = run_tesseract_ocr(uploaded_file.storage_path, language=language)
        metadata = {
            **result.metadata,
            "filename": uploaded_file.original_name,
            "file_id": uploaded_file.id,
        }
        return ParsedDocument(text=result.text, metadata=metadata)


DOCUMENT_PARSERS: tuple[DocumentParser, ...] = (
    TxtDocumentParser(),
    CsvDocumentParser(),
    JsonDocumentParser(),
    DocxDocumentParser(),
    PdfDocumentParser(),
)


def get_document_parser(extension: str, strategy: str = "auto") -> DocumentParser:
    if strategy == "ocr":
        return OcrDocumentParser()

    for parser in DOCUMENT_PARSERS:
        if parser.supports(extension):
            return parser

    raise ValueError(f"No parser is available for extension '{extension}'.")


def parse_uploaded_file(uploaded_file: UploadedFile, strategy: str = "auto") -> ParsedDocument:
    parser = get_document_parser(uploaded_file.extension, strategy=strategy)
    return parser.parse(uploaded_file)
